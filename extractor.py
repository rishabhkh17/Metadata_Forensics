# # # #extracts metadata from files (images, videos, PDFs, docx)

# # # from PIL import Image
# # # from PIL import Image, ExifTags
# # # from PIL.ExifTags import TAGS
# # # import os
# # # import PyPDF2
# # # import docx
# # # from pymediainfo import MediaInfo

# # # def convert_metadata(metadata):
# # #     """Converts IFDRational and other non-serializable types to JSON-friendly format."""
# # #     converted = {}
# # #     for key, value in metadata.items():
# # #         if isinstance(value, bytes):  # Convert binary data to string
# # #             converted[key] = value.decode(errors="ignore")
# # #         elif isinstance(value, tuple):  # Convert tuples to lists
# # #             converted[key] = list(value)
# # #         elif isinstance(value, int) or isinstance(value, float) or isinstance(value, str):
# # #             converted[key] = value
# # #         else:  # Convert IFDRational and other non-serializable objects
# # #             converted[key] = str(value)
# # #     return converted

# # # def extract_image_metadata(file_path):
# # #     """Extract metadata from image files (JPEG, PNG, etc.)."""
# # #     metadata = {}
# # #     try:
# # #         img = Image.open(file_path)

# # #         exif_data = img.getexif()

# # #         if exif_data:
# # #             for tag, value in exif_data.items():
# # #                 tag_name = TAGS.get(tag, tag)
# # #                 metadata[tag_name] = value
        
# # #         # Convert metadata to a JSON-friendly format
# # #         return convert_metadata(metadata)

# # #     except Exception as e:
# # #         return {"Error": f"Failed to extract image metadata: {str(e)}"}

# # # def extract_pdf_metadata(file_path):
# # #     """Extract metadata from a PDF file."""
# # #     metadata = {}
# # #     with open(file_path, "rb") as pdf_file:
# # #         reader = PyPDF2.PdfReader(pdf_file)
# # #         info = reader.metadata
# # #         if info:
# # #             metadata = {key.lstrip("/"): info[key] for key in info}
# # #     return metadata

# # # def extract_docx_metadata(file_path):
# # #     """Extract metadata from a Word document."""
# # #     metadata = {}
# # #     doc = docx.Document(file_path)
# # #     core_props = doc.core_properties
# # #     metadata["author"] = core_props.author or "Unknown"
# # #     metadata["created"] = core_props.created or "Unknown"
# # #     metadata["modified"] = core_props.modified or "Unknown"
# # #     return metadata


# # # def extract_video_metadata(file_path):
# # #     """Extract metadata from a video file using pymediainfo."""
# # #     metadata = {}
# # #     media_info = MediaInfo.parse(file_path)

# # #     for track in media_info.tracks:
# # #         if track.track_type == "General":
# # #             metadata["Duration"] = track.duration
# # #             metadata["File size"] = track.file_size
# # #             metadata["Format"] = track.format
# # #             metadata["Bitrate"] = track.overall_bit_rate
# # #         elif track.track_type == "Video":
# # #             metadata["Width"] = track.width
# # #             metadata["Height"] = track.height
# # #             metadata["Frame rate"] = track.frame_rate
# # #             metadata["Codec"] = track.codec_id
    
# # #     return metadata


# # # def extract_metadata(file_path):
# # #     """Determines file type and extracts metadata accordingly."""
# # #     _, ext = os.path.splitext(file_path)
# # #     ext = ext.lower()

# # #     if ext in [".jpg", ".jpeg", ".png"]:
# # #         return extract_image_metadata(file_path)
# # #     elif ext == ".pdf":
# # #         return extract_pdf_metadata(file_path)
# # #     elif ext == ".docx":
# # #         return extract_docx_metadata(file_path)
# # #     elif ext in [".mp4", ".avi", ".mkv"]:
# # #         return extract_video_metadata(file_path)
# # #     else:
# # #         return {"Error": "Unsupported file type"}


# # # Extracts metadata from files (images, videos, PDFs, docx)
# # from PIL import Image
# # from PIL.ExifTags import TAGS
# # import os
# # import PyPDF2
# # import docx
# # from pymediainfo import MediaInfo

# # def convert_metadata(metadata):
# #     """Converts IFDRational and other non-serializable types to JSON-friendly format."""
# #     converted = {}
# #     for key, value in metadata.items():
# #         if isinstance(value, bytes):  # Convert binary data to string
# #             converted[key] = value.decode(errors="ignore")
# #         elif isinstance(value, tuple):  # Convert tuples to lists
# #             converted[key] = list(value)
# #         elif isinstance(value, int) or isinstance(value, float) or isinstance(value, str):
# #             converted[key] = value
# #         else:  # Convert IFDRational and other non-serializable objects
# #             converted[key] = str(value)
# #     return converted

# # def extract_image_metadata(file_path):
# #     """Extract metadata from image files (JPEG, PNG, etc.)."""
# #     metadata = {}
# #     try:
# #         img = Image.open(file_path)
        
# #         # Basic image info
# #         metadata["format"] = img.format
# #         metadata["mode"] = img.mode
# #         metadata["size"] = img.size
        
# #         # Try multiple methods to get EXIF data
# #         exif_data = img.getexif()
        
# #         # If getexif() returns None or empty, try the private method as fallback
# #         if not exif_data:
# #             try:
# #                 exif_data = img._getexif() or {}
# #             except (AttributeError, TypeError):
# #                 exif_data = {}
        
# #         # Add all EXIF data to metadata
# #         for tag, value in exif_data.items():
# #             tag_name = TAGS.get(tag, tag)
# #             metadata[tag_name] = value
        
# #         # Get all image info
# #         img_info = img.info
# #         for key, value in img_info.items():
# #             if key not in metadata:
# #                 metadata[key] = value
                
# #         # Convert metadata to a JSON-friendly format
# #         return convert_metadata(metadata)
# #     except Exception as e:
# #         return {"Error": f"Failed to extract image metadata: {str(e)}"}

# # def extract_pdf_metadata(file_path):
# #     """Extract metadata from a PDF file."""
# #     metadata = {}
# #     try:
# #         with open(file_path, "rb") as pdf_file:
# #             reader = PyPDF2.PdfReader(pdf_file)
# #             info = reader.metadata
# #             if info:
# #                 metadata = {key.lstrip("/"): info[key] for key in info}
# #             metadata["Pages"] = len(reader.pages)
# #         return metadata
# #     except Exception as e:
# #         return {"Error": f"Failed to extract PDF metadata: {str(e)}"}

# # def extract_docx_metadata(file_path):
# #     """Extract metadata from a Word document."""
# #     metadata = {}
# #     try:
# #         doc = docx.Document(file_path)
# #         core_props = doc.core_properties
# #         metadata["author"] = core_props.author or "Unknown"
# #         metadata["created"] = core_props.created or "Unknown"
# #         metadata["modified"] = core_props.modified or "Unknown"
# #         metadata["title"] = core_props.title or "Unknown"
# #         metadata["last_modified_by"] = core_props.last_modified_by or "Unknown"
# #         return metadata
# #     except Exception as e:
# #         return {"Error": f"Failed to extract DOCX metadata: {str(e)}"}

# # def extract_video_metadata(file_path):
# #     """Extract metadata from a video file using pymediainfo."""
# #     metadata = {}
# #     try:
# #         media_info = MediaInfo.parse(file_path)
# #         for track in media_info.tracks:
# #             if track.track_type == "General":
# #                 metadata["Duration"] = track.duration
# #                 metadata["File size"] = track.file_size
# #                 metadata["Format"] = track.format
# #                 metadata["Bitrate"] = track.overall_bit_rate
# #             elif track.track_type == "Video":
# #                 metadata["Width"] = track.width
# #                 metadata["Height"] = track.height
# #                 metadata["Frame rate"] = track.frame_rate
# #                 metadata["Codec"] = track.codec_id
# #         return metadata
# #     except Exception as e:
# #         return {"Error": f"Failed to extract video metadata: {str(e)}"}

# # def extract_metadata(file_path):
# #     """Determines file type and extracts metadata accordingly."""
# #     _, ext = os.path.splitext(file_path)  # Fixed syntax error here
# #     ext = ext.lower()
    
# #     if ext in [".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".gif"]:
# #         return extract_image_metadata(file_path)
# #     elif ext == ".pdf":
# #         return extract_pdf_metadata(file_path)
# #     elif ext == ".docx":
# #         return extract_docx_metadata(file_path)
# #     elif ext in [".mp4", ".avi", ".mkv", ".mov", ".wmv"]:
# #         return extract_video_metadata(file_path)
# #     else:
# #         return {"Error": "Unsupported file type"}


# # comprehensive_metadata_extractor.py
# # A tool to extract detailed metadata from various file types

# import os
# import sys
# import json
# from datetime import datetime
# from PIL import Image
# from PIL.ExifTags import TAGS, GPSTAGS
# import PyPDF2
# import docx
# from pymediainfo import MediaInfo

# def convert_metadata(metadata):
#     """Converts complex data types to JSON-friendly format."""
#     converted = {}
#     for key, value in metadata.items():
#         if isinstance(value, bytes):
#             # Convert binary data to string
#             converted[key] = value.decode(errors="ignore")
#         elif isinstance(value, (datetime, tuple, list, dict)):
#             # Convert complex types to string
#             converted[key] = str(value)
#         elif isinstance(value, int) or isinstance(value, float) or isinstance(value, str):
#             # Pass through primitive types
#             converted[key] = value
#         else:
#             # Convert other types (like IFDRational) to string
#             converted[key] = str(value)
#     return converted

# def extract_image_metadata(file_path):
#     """
#     Extract detailed metadata from image files (JPEG, PNG, TIFF, etc.)
#     Includes EXIF, ICC profile, and format-specific information.
#     """
#     metadata = {}
#     try:
#         img = Image.open(file_path)
        
#         # Basic image properties
#         metadata["file_name"] = os.path.basename(file_path)
#         metadata["file_size"] = os.path.getsize(file_path)
#         metadata["format"] = img.format
#         metadata["mode"] = img.mode
#         metadata["width"] = img.width
#         metadata["height"] = img.height
#         metadata["color_profile"] = {
#             "1": "1-bit pixels (black and white)",
#             "L": "8-bit grayscale",
#             "P": "8-bit palette-mapped",
#             "RGB": "3x8-bit RGB",
#             "RGBA": "4x8-bit RGBA",
#             "CMYK": "4x8-bit CMYK",
#             "YCbCr": "3x8-bit YCbCr",
#             "I": "32-bit integer pixels",
#             "F": "32-bit float pixels"
#         }.get(img.mode, img.mode)
        
#         # Try multiple methods to get EXIF data
#         exif_data = {}
        
#         # Method 1: Public getexif() method
#         try:
#             exif = img.getexif()
#             if exif:
#                 for tag, value in exif.items():
#                     tag_name = TAGS.get(tag, tag)
#                     exif_data[tag_name] = value
#         except Exception as e:
#             pass
            
#         # Method 2: Private _getexif() method (fallback)
#         if not exif_data:
#             try:
#                 exif = img._getexif()
#                 if exif:
#                     for tag, value in exif.items():
#                         tag_name = TAGS.get(tag, tag)
#                         exif_data[tag_name] = value
#             except Exception as e:
#                 pass
        
#         # Common EXIF tags with human-readable names
#         common_tags = {
#             "Make": "Camera Make",
#             "Model": "Camera Model",
#             "DateTime": "Date/Time",
#             "DateTimeOriginal": "Original Date/Time",
#             "ExposureTime": "Exposure Time",
#             "FNumber": "F-Number",
#             "ISOSpeedRatings": "ISO Speed",
#             "FocalLength": "Focal Length",
#             "Flash": "Flash Used",
#             "XResolution": "X Resolution",
#             "YResolution": "Y Resolution",
#             "Software": "Software Used",
#             "Artist": "Artist",
#             "Copyright": "Copyright"
#         }
        
#         # Extract key EXIF data with better labels
#         for tag, readable_name in common_tags.items():
#             if tag in exif_data:
#                 metadata[readable_name] = exif_data[tag]
        
#         # Extract GPS info if available
#         try:
#             if "GPSInfo" in exif_data:
#                 gps_info = {}
#                 for key, value in exif_data["GPSInfo"].items():
#                     tag_name = GPSTAGS.get(key, key)
#                     gps_info[tag_name] = value
                
#                 # Calculate decimal coordinates if possible
#                 if all(x in gps_info for x in ["GPSLatitude", "GPSLatitudeRef", "GPSLongitude", "GPSLongitudeRef"]):
#                     try:
#                         lat = gps_info["GPSLatitude"]
#                         lat_ref = gps_info["GPSLatitudeRef"]
#                         lon = gps_info["GPSLongitude"]
#                         lon_ref = gps_info["GPSLongitudeRef"]
                        
#                         lat_decimal = lat[0] + lat[1]/60 + lat[2]/3600
#                         if lat_ref == "S":
#                             lat_decimal = -lat_decimal
                            
#                         lon_decimal = lon[0] + lon[1]/60 + lon[2]/3600
#                         if lon_ref == "W":
#                             lon_decimal = -lon_decimal
                            
#                         gps_info["Decimal_Coordinates"] = f"{lat_decimal}, {lon_decimal}"
#                     except:
#                         pass
                
#                 metadata["GPS_Information"] = gps_info
#         except Exception as e:
#            pass
        
#         # ICC profile info
#         if "icc_profile" in img.info:
#             metadata["has_icc_profile"] = True
        
#         # Format-specific metadata
#         if img.format == "PNG":
#             # PNG-specific chunks
#             try:
#                 # Store PNG chunks types
#                 chunks = [chunk[0].decode('ascii') for chunk in img.png.chunks 
#                           if chunk[0].decode('ascii') not in ["IDAT", "IHDR"]]
#                 if chunks:
#                     metadata["PNG_chunks"] = chunks
#             except Exception as e:
#                 pass
        
#         # JPEG specific
#         if img.format == "JPEG":
#             # Check for IPTC data
#             if "iptc" in img.info:
#                 metadata["has_iptc_data"] = True
            
#             # Check for Adobe XMP data
#             if "xml:com.adobe.xmp" in img.info:
#                 metadata["has_xmp_data"] = True
        
#         # TIFF specific
#         if img.format == "TIFF":
#             if hasattr(img, "tag"):
#                 metadata["tiff_tags_count"] = len(img.tag.items())
        
#         # Include any other relevant data from img.info
#         for key, value in img.info.items():
#             if isinstance(key, str) and key not in ["exif", "icc_profile"] and not key.startswith("xml:"):
#                 try:
#                     if isinstance(value, (int, float, str)):
#                         metadata[f"info_{key}"] = value
#                     else:
#                         metadata[f"info_{key}_present"] = True
#                 except:
#                     pass
                    
#         return convert_metadata(metadata)
    
#     except Exception as e:
#         return {"Error": f"Failed to extract image metadata: {str(e)}"}

# def extract_pdf_metadata(file_path):
#     """
#     Extract detailed metadata from PDF files.
#     Handles different PyPDF2 versions.
#     """
#     metadata = {}
#     try:
#         with open(file_path, "rb") as pdf_file:
#             metadata["file_name"] = os.path.basename(file_path)
#             metadata["file_size"] = os.path.getsize(file_path)
            
#             reader = PyPDF2.PdfReader(pdf_file)
            
#             # Basic document information
#             metadata["page_count"] = len(reader.pages)
            
#             # Handle different PyPDF2 versions for pdf_version attribute
#             try:
#                 # Try different possible attribute names for version
#                 if hasattr(reader, "pdf_version"):
#                     metadata["pdf_version"] = f"{reader.pdf_version}"
#                 elif hasattr(reader, "version"):
#                     metadata["pdf_version"] = f"{reader.version}"
#                 elif hasattr(reader, "_version"):
#                     metadata["pdf_version"] = f"{reader._version}"
#                 else:
#                     # Check the trailer dictionary
#                     if hasattr(reader, "trailer") and reader.trailer and "/Root" in reader.trailer:
#                         root = reader.trailer["/Root"]
#                         if "/Version" in root:
#                             metadata["pdf_version"] = str(root["/Version"])
#                         else:
#                             metadata["pdf_version"] = "Unknown"
#                     else:
#                         metadata["pdf_version"] = "Unknown"
#             except Exception as e:
#                 metadata["pdf_version"] = "Unknown"
#                 metadata["version_error"] = str(e)
            
#             # Document information dictionary
#             info = reader.metadata
#             if info:
#                 # Standard metadata fields
#                 for field in dir(info):
#                     if not field.startswith("_") and not callable(getattr(info, field)):
#                         value = getattr(info, field)
#                         if value is not None:
#                             metadata[field] = str(value)
            
#             # Sample page sizes (first 3 pages)
#             page_sizes = []
#             for i in range(min(3, len(reader.pages))):
#                 page = reader.pages[i]
#                 if "/MediaBox" in page:
#                     page_sizes.append(str(page["/MediaBox"]))
#             if page_sizes:
#                 metadata["sample_page_sizes"] = page_sizes
            
#             # Security information - handle version differences
#             try:
#                 if hasattr(reader, "is_encrypted"):
#                     metadata["is_encrypted"] = reader.is_encrypted
#                 else:
#                     # Fall back for older versions
#                     metadata["is_encrypted"] = "Unknown"
#             except Exception as e:
#                 metadata["is_encrypted"] = "Unknown"
#                 metadata["encryption_check_error"] = str(e)
            
#             # Check for XMP metadata
#             try:
#                 if hasattr(reader, "xmp_metadata") and reader.xmp_metadata:
#                     metadata["has_xmp_metadata"] = True
#             except Exception as e:
#                 metadata["xmp_check_error"] = str(e)
                
#             # Check for form fields - handle version differences
#             try:
#                 has_fields = False
#                 # Try different methods to get fields
#                 if hasattr(reader, "get_fields"):
#                     fields = reader.get_fields()
#                     has_fields = fields is not None and len(fields) > 0
#                 elif hasattr(reader, "getFields"):
#                     fields = reader.getFields()
#                     has_fields = fields is not None and len(fields) > 0
#                 elif hasattr(reader, "getFormTextFields"):
#                     fields = reader.getFormTextFields()
#                     has_fields = fields is not None and len(fields) > 0
                
#                 if has_fields:
#                     metadata["has_form_fields"] = True
#                     metadata["form_field_count"] = len(fields)
#             except Exception as e:
#                 metadata["form_check_error"] = str(e)
                
#             # Check for annotations
#             try:
#                 annotation_count = 0
#                 for i in range(min(5, len(reader.pages))):  # Just check first 5 pages
#                     page = reader.pages[i]
#                     if "/Annots" in page:
#                         annots = page["/Annots"]
#                         if annots:
#                             annotation_count += len(annots)
#                 if annotation_count > 0:
#                     metadata["has_annotations"] = True
#                     metadata["annotation_count_sample"] = annotation_count
#             except Exception as e:
#                 metadata["annotation_check_error"] = str(e)
                
#         return convert_metadata(metadata)
    
#     except Exception as e:
#         return {"Error": f"Failed to extract PDF metadata: {str(e)}"}
    
    
# def extract_docx_metadata(file_path):
#     """
#     Extract detailed metadata from Word documents.
#     Includes document properties, statistics, and content information.
#     """
#     metadata = {}
#     try:
#         metadata["file_name"] = os.path.basename(file_path)
#         metadata["file_size"] = os.path.getsize(file_path)
        
#         doc = docx.Document(file_path)
        
#         # Core properties
#         core_props = doc.core_properties
#         metadata["title"] = core_props.title or "Unknown"
#         metadata["author"] = core_props.author or "Unknown"
#         metadata["created"] = str(core_props.created) if core_props.created else "Unknown"
#         metadata["modified"] = str(core_props.modified) if core_props.modified else "Unknown"
#         metadata["last_modified_by"] = core_props.last_modified_by or "Unknown"
        
#         if core_props.revision:
#             metadata["revision"] = core_props.revision
#         if core_props.category:
#             metadata["category"] = core_props.category
#         if core_props.comments:
#             metadata["comments"] = core_props.comments
#         if core_props.subject:
#             metadata["subject"] = core_props.subject
#         if core_props.keywords:
#             metadata["keywords"] = core_props.keywords
#         if core_props.language:
#             metadata["language"] = core_props.language
#         if hasattr(core_props, 'content_status') and core_props.content_status:
#             metadata["content_status"] = core_props.content_status
            
#         # Document statistics
#         metadata["paragraph_count"] = len(doc.paragraphs)
#         metadata["section_count"] = len(doc.sections)
#         metadata["page_count"] = len(doc.sections)  # Approximate
        
#         # Table information
#         tables = doc.tables
#         metadata["table_count"] = len(tables)
        
#         # Style information
#         styles = doc.styles
#         metadata["style_count"] = len(styles)
        
#         # Header/footer information
#         has_header = False
#         has_footer = False
#         for section in doc.sections:
#             if section.header.is_linked_to_previous == False:
#                 has_header = True
#             if section.footer.is_linked_to_previous == False:
#                 has_footer = True
        
#         metadata["has_headers"] = has_header
#         metadata["has_footers"] = has_footer
        
#         # Check for images
#         try:
#             image_count = 0
#             for rel in doc.part.rels.values():
#                 if "image" in rel.target_ref:
#                     image_count += 1
#             metadata["image_count"] = image_count
#         except Exception as e:
#            pass
            
#         return convert_metadata(metadata)
    
#     except Exception as e:
#         return {"Error": f"Failed to extract DOCX metadata: {str(e)}"}

# def extract_video_metadata(file_path):
#     """
#     Extract detailed metadata from video files using pymediainfo.
#     Includes video, audio, and general properties.
#     """
#     metadata = {}
#     try:
#         metadata["file_name"] = os.path.basename(file_path)
#         metadata["file_size"] = os.path.getsize(file_path)
        
#         media_info = MediaInfo.parse(file_path)
        
#         # General information
#         for track in media_info.tracks:
#             if track.track_type == "General":
#                 general_info = {
#                     "Format": track.format,
#                     "Format Profile": track.format_profile if hasattr(track, 'format_profile') else None,
#                     "Codec": track.codec_id if hasattr(track, 'codec_id') else None,
#                     "Duration": f"{track.duration} ms" if hasattr(track, 'duration') else None,
#                     "Duration_Formatted": track.other_duration[0] if hasattr(track, 'other_duration') else None,
#                     "File Size": f"{track.file_size} bytes" if hasattr(track, 'file_size') else None,
#                     "Overall Bitrate": f"{track.overall_bit_rate} bps" if hasattr(track, 'overall_bit_rate') else None,
#                     "Created Date": track.file_created_date if hasattr(track, 'file_created_date') else None,
#                     "Modified Date": track.file_modified_date if hasattr(track, 'file_modified_date') else None,
#                     "Encoded by": track.encoded_by if hasattr(track, 'encoded_by') else None,
#                     "Writing Application": track.writing_application if hasattr(track, 'writing_application') else None,
#                 }
#                 # Filter out None values
#                 metadata["General"] = {k: v for k, v in general_info.items() if v is not None}
                
#             # Video information
#             elif track.track_type == "Video":
#                 video_info = {
#                     "Format": track.format,
#                     "Format Profile": track.format_profile if hasattr(track, 'format_profile') else None,
#                     "Codec": track.codec_id if hasattr(track, 'codec_id') else None,
#                     "Width": f"{track.width} pixels" if hasattr(track, 'width') else None,
#                     "Height": f"{track.height} pixels" if hasattr(track, 'height') else None,
#                     "Display Aspect Ratio": track.display_aspect_ratio if hasattr(track, 'display_aspect_ratio') else None,
#                     "Frame Rate": f"{track.frame_rate} fps" if hasattr(track, 'frame_rate') else None,
#                     "Bit Depth": f"{track.bit_depth} bits" if hasattr(track, 'bit_depth') else None,
#                     "Bit Rate": f"{track.bit_rate} bps" if hasattr(track, 'bit_rate') else None,
#                     "Scan Type": track.scan_type if hasattr(track, 'scan_type') else None,
#                     "Color Space": track.color_space if hasattr(track, 'color_space') else None,
#                     "Chroma Subsampling": track.chroma_subsampling if hasattr(track, 'chroma_subsampling') else None,
#                 }
#                 # Filter out None values
#                 metadata["Video"] = {k: v for k, v in video_info.items() if v is not None}
                
#             # Audio information
#             elif track.track_type == "Audio":
#                 audio_info = {
#                     "Format": track.format,
#                     "Format Profile": track.format_profile if hasattr(track, 'format_profile') else None,
#                     "Codec": track.codec_id if hasattr(track, 'codec_id') else None,
#                     "Channel(s)": track.channel_s if hasattr(track, 'channel_s') else None,
#                     "Sampling Rate": f"{track.sampling_rate} Hz" if hasattr(track, 'sampling_rate') else None,
#                     "Bit Rate": f"{track.bit_rate} bps" if hasattr(track, 'bit_rate') else None,
#                     "Compression Mode": track.compression_mode if hasattr(track, 'compression_mode') else None,
#                     "Language": track.language if hasattr(track, 'language') else None,
#                     "Default": track.default if hasattr(track, 'default') else None,
#                 }
#                 # Filter out None values
#                 if "Audio" not in metadata:
#                     metadata["Audio"] = []
#                 metadata["Audio"].append({k: v for k, v in audio_info.items() if v is not None})
                
#             # Subtitle information
#             elif track.track_type == "Text":
#                 text_info = {
#                     "Format": track.format,
#                     "Codec ID": track.codec_id if hasattr(track, 'codec_id') else None,
#                     "Language": track.language if hasattr(track, 'language') else None,
#                     "Default": track.default if hasattr(track, 'default') else None,
#                     "Forced": track.forced if hasattr(track, 'forced') else None,
#                 }
#                 # Filter out None values
#                 if "Subtitles" not in metadata:
#                     metadata["Subtitles"] = []
#                 metadata["Subtitles"].append({k: v for k, v in text_info.items() if v is not None})
        
#         return convert_metadata(metadata)
    
#     except Exception as e:
#         return {"Error": f"Failed to extract video metadata: {str(e)}"}

# def extract_metadata(file_path):
#     """
#     Determine file type and extract appropriate metadata.
#     Returns a dictionary with detailed metadata based on file type.
#     """
#     if not os.path.exists(file_path):
#         return {"Error": f"File not found: {file_path}"}
    
#     _, ext = os.path.splitext(file_path)
#     ext = ext.lower()
    
#     # Process based on file extension
#     if ext in [".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"]:
#         return extract_image_metadata(file_path)
#     elif ext == ".pdf":
#         return extract_pdf_metadata(file_path)
#     elif ext == ".docx":
#         return extract_docx_metadata(file_path)
#     elif ext in [".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm"]:
#         return extract_video_metadata(file_path)
#     else:
#         return {"Error": f"Unsupported file type: {ext}"}

# def main():
#     """
#     Main function: demonstrate usage with a sample file.
#     """
#     if len(sys.argv) != 2:
#         print("Usage: python metadata_extractor.py <file_path>")
#         sys.exit(1)
    
#     file_path = sys.argv[1]
#     metadata = extract_metadata(file_path)
    
#     # Print formatted JSON output
#     print(json.dumps(metadata, indent=2))

# if __name__ == "__main__":
#     main()

# comprehensive_metadata_extractor.py
# A tool to extract detailed metadata from various file types

import os
import sys
import json
from datetime import datetime
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import PyPDF2
import docx
from pymediainfo import MediaInfo

def convert_metadata(metadata):
    """Converts complex data types to JSON-friendly format."""
    converted = {}
    for key, value in metadata.items():
        if isinstance(value, bytes):
            # Convert binary data to string
            converted[key] = value.decode(errors="ignore")
        elif isinstance(value, (datetime, tuple, list, dict)):
            # Convert complex types to string
            converted[key] = str(value)
        elif isinstance(value, int) or isinstance(value, float) or isinstance(value, str):
            # Pass through primitive types
            converted[key] = value
        else:
            # Convert other types (like IFDRational) to string
            converted[key] = str(value)
    return converted

def parse_pdf_date(date_str):
    """
    Parse PDF date format which can be in multiple different formats.
    Handles format with or without timezone information.
    """
    if not date_str:
        return None
        
    try:
        # Remove 'D:' prefix if present
        if date_str.startswith('D:'):
            date_str = date_str[2:]
            
        # Basic case: YYYYMMDDHHmmSS
        if len(date_str) == 14 and date_str.isdigit():
            return f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]} {date_str[8:10]}:{date_str[10:12]}:{date_str[12:14]}"
            
        # Case with timezone: YYYYMMDDHHmmSS+HH'mm'
        if len(date_str) > 14 and date_str[:14].isdigit():
            time_part = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]} {date_str[8:10]}:{date_str[10:12]}:{date_str[12:14]}"
            # Add timezone info if available
            timezone_part = date_str[14:].replace("'", "")
            if timezone_part:
                time_part += f" {timezone_part}"
            return time_part
            
        # If it doesn't match expected formats, return as is
        return date_str
    except Exception:
        # If parsing fails, return the original string
        return date_str

def extract_image_metadata(file_path):
    """
    Extract detailed metadata from image files (JPEG, PNG, TIFF, etc.)
    Includes EXIF, ICC profile, and format-specific information.
    """
    metadata = {}
    try:
        img = Image.open(file_path)
        
        # Basic image properties
        metadata["file_name"] = os.path.basename(file_path)
        metadata["file_size"] = os.path.getsize(file_path)
        metadata["format"] = img.format
        metadata["mode"] = img.mode
        metadata["width"] = img.width
        metadata["height"] = img.height
        metadata["color_profile"] = {
            "1": "1-bit pixels (black and white)",
            "L": "8-bit grayscale",
            "P": "8-bit palette-mapped",
            "RGB": "3x8-bit RGB",
            "RGBA": "4x8-bit RGBA",
            "CMYK": "4x8-bit CMYK",
            "YCbCr": "3x8-bit YCbCr",
            "I": "32-bit integer pixels",
            "F": "32-bit float pixels"
        }.get(img.mode, img.mode)
        
        # Try multiple methods to get EXIF data
        exif_data = {}
        
        # Method 1: Public getexif() method
        try:
            exif = img.getexif()
            if exif:
                for tag, value in exif.items():
                    tag_name = TAGS.get(tag, tag)
                    exif_data[tag_name] = value
        except Exception as e:
            print(f"Error getting EXIF data (method 1): {e}")
            
        # Method 2: Private _getexif() method (fallback)
        if not exif_data:
            try:
                exif = img._getexif()
                if exif:
                    for tag, value in exif.items():
                        tag_name = TAGS.get(tag, tag)
                        exif_data[tag_name] = value
            except Exception as e:
                print(f"Error getting EXIF data (method 2): {e}")
        
        # Common EXIF tags with human-readable names
        common_tags = {
            "Make": "Camera Make",
            "Model": "Camera Model",
            "DateTime": "Date/Time",
            "DateTimeOriginal": "Original Date/Time",
            "ExposureTime": "Exposure Time",
            "FNumber": "F-Number",
            "ISOSpeedRatings": "ISO Speed",
            "FocalLength": "Focal Length",
            "Flash": "Flash Used",
            "XResolution": "X Resolution",
            "YResolution": "Y Resolution",
            "Software": "Software Used",
            "Artist": "Artist",
            "Copyright": "Copyright"
        }
        
        # Extract key EXIF data with better labels
        for tag, readable_name in common_tags.items():
            if tag in exif_data:
                metadata[readable_name] = exif_data[tag]
        
        # Extract GPS info if available
        try:
            if "GPSInfo" in exif_data:
                gps_info = {}
                for key, value in exif_data["GPSInfo"].items():
                    tag_name = GPSTAGS.get(key, key)
                    gps_info[tag_name] = value
                
                # Calculate decimal coordinates if possible
                if all(x in gps_info for x in ["GPSLatitude", "GPSLatitudeRef", "GPSLongitude", "GPSLongitudeRef"]):
                    try:
                        lat = gps_info["GPSLatitude"]
                        lat_ref = gps_info["GPSLatitudeRef"]
                        lon = gps_info["GPSLongitude"]
                        lon_ref = gps_info["GPSLongitudeRef"]
                        
                        lat_decimal = lat[0] + lat[1]/60 + lat[2]/3600
                        if lat_ref == "S":
                            lat_decimal = -lat_decimal
                            
                        lon_decimal = lon[0] + lon[1]/60 + lon[2]/3600
                        if lon_ref == "W":
                            lon_decimal = -lon_decimal
                            
                        gps_info["Decimal_Coordinates"] = f"{lat_decimal}, {lon_decimal}"
                    except Exception as e:
                        print(f"Error calculating GPS coordinates: {e}")
                
                metadata["GPS_Information"] = gps_info
        except Exception as e:
            print(f"Error extracting GPS data: {e}")
        
        # ICC profile info
        if "icc_profile" in img.info:
            metadata["has_icc_profile"] = True
        
        # Format-specific metadata
        if img.format == "PNG":
            # PNG-specific chunks
            try:
                # Store PNG chunks types
                chunks = [chunk[0].decode('ascii') for chunk in img.png.chunks 
                          if chunk[0].decode('ascii') not in ["IDAT", "IHDR"]]
                if chunks:
                    metadata["PNG_chunks"] = chunks
            except Exception as e:
                print(f"Error extracting PNG chunks: {e}")
        
        # JPEG specific
        if img.format == "JPEG":
            # Check for IPTC data
            if "iptc" in img.info:
                metadata["has_iptc_data"] = True
            
            # Check for Adobe XMP data
            if "xml:com.adobe.xmp" in img.info:
                metadata["has_xmp_data"] = True
        
        # TIFF specific
        if img.format == "TIFF":
            if hasattr(img, "tag"):
                metadata["tiff_tags_count"] = len(img.tag.items())
        
        # Include any other relevant data from img.info
        for key, value in img.info.items():
            if isinstance(key, str) and key not in ["exif", "icc_profile"] and not key.startswith("xml:"):
                try:
                    if isinstance(value, (int, float, str)):
                        metadata[f"info_{key}"] = value
                    else:
                        metadata[f"info_{key}_present"] = True
                except Exception:
                    pass
                    
        return convert_metadata(metadata)
    
    except Exception as e:
        print(f"Failed to extract image metadata: {e}")
        return {"Error": "Failed to extract image metadata"}

def extract_pdf_metadata(file_path):
    """
    Extract detailed metadata from PDF files.
    Includes document properties, security settings, and structure information.
    """
    metadata = {}
    try:
        with open(file_path, "rb") as pdf_file:
            metadata["file_name"] = os.path.basename(file_path)
            metadata["file_size"] = os.path.getsize(file_path)
            
            reader = PyPDF2.PdfReader(pdf_file)
            
            # Basic document information
            metadata["page_count"] = len(reader.pages)
            
            # Handle different PyPDF2 versions for pdf_version attribute
            try:
                # Try different possible attribute names for version
                if hasattr(reader, "pdf_version"):
                    metadata["pdf_version"] = f"{reader.pdf_version}"
                elif hasattr(reader, "version"):
                    metadata["pdf_version"] = f"{reader.version}"
                elif hasattr(reader, "_version"):
                    metadata["pdf_version"] = f"{reader._version}"
                else:
                    # Version not directly accessible, set to unknown
                    metadata["pdf_version"] = "Unknown"
            except Exception as e:
                print(f"Error getting PDF version: {e}")
                metadata["pdf_version"] = "Unknown"
            
            # Document information dictionary - handles date formats
            try:
                info = reader.metadata
                if info:
                    for name in dir(info):
                        if not name.startswith("_") and not callable(getattr(info, name)):
                            value = getattr(info, name)
                            
                            # Handle date fields differently
                            if name in ['creation_date', 'modification_date', 'created', 'modified'] and value:
                                try:
                                    # Try to parse PDF date format
                                    parsed_date = parse_pdf_date(str(value))
                                    metadata[name] = parsed_date
                                except Exception as e:
                                    print(f"Error parsing PDF date: {e}")
                                    # If parsing fails, just use the string value
                                    metadata[name] = str(value)
                            elif value is not None:
                                metadata[name] = str(value)
            except Exception as e:
                print(f"Error extracting PDF metadata: {e}")
            
            # Sample page sizes
            try:
                page_sizes = []
                for i in range(min(3, len(reader.pages))):
                    page = reader.pages[i]
                    if "/MediaBox" in page:
                        page_sizes.append(str(page["/MediaBox"]))
                if page_sizes:
                    metadata["sample_page_sizes"] = page_sizes
            except Exception as e:
                print(f"Error getting page sizes: {e}")
            
            # Security information
            try:
                if hasattr(reader, "is_encrypted"):
                    metadata["is_encrypted"] = reader.is_encrypted
                else:
                    metadata["is_encrypted"] = "Unknown"
            except Exception as e:
                print(f"Error checking encryption: {e}")
                
            # Check for form fields
            try:
                has_fields = False
                fields = None
                
                # Try different methods to get fields
                methods = ["get_fields", "getFields", "getFormTextFields"]
                for method_name in methods:
                    if hasattr(reader, method_name):
                        method = getattr(reader, method_name)
                        try:
                            fields = method()
                            if fields:
                                has_fields = True
                                break
                        except Exception:
                            continue
                
                if has_fields and fields:
                    metadata["has_form_fields"] = True
                    metadata["form_field_count"] = len(fields)
            except Exception as e:
                print(f"Error checking form fields: {e}")
            
            # Check for annotations
            try:
                annotation_count = 0
                for i in range(min(5, len(reader.pages))):
                    page = reader.pages[i]
                    if "/Annots" in page:
                        annots = page["/Annots"]
                        if annots:
                            annotation_count += len(annots)
                if annotation_count > 0:
                    metadata["has_annotations"] = True
                    metadata["annotation_count_sample"] = annotation_count
            except Exception as e:
                print(f"Error checking annotations: {e}")
                
        return convert_metadata(metadata)
    
    except Exception as e:
        print(f"Failed to extract PDF metadata: {e}")
        return {"Error": "Failed to extract PDF metadata"}

def extract_docx_metadata(file_path):
    """
    Extract detailed metadata from Word documents.
    Includes document properties, statistics, and content information.
    """
    metadata = {}
    try:
        metadata["file_name"] = os.path.basename(file_path)
        metadata["file_size"] = os.path.getsize(file_path)
        
        doc = docx.Document(file_path)
        
        # Core properties
        core_props = doc.core_properties
        metadata["title"] = core_props.title or "Unknown"
        metadata["author"] = core_props.author or "Unknown"
        metadata["created"] = str(core_props.created) if core_props.created else "Unknown"
        metadata["modified"] = str(core_props.modified) if core_props.modified else "Unknown"
        metadata["last_modified_by"] = core_props.last_modified_by or "Unknown"
        
        if core_props.revision:
            metadata["revision"] = core_props.revision
        if core_props.category:
            metadata["category"] = core_props.category
        if core_props.comments:
            metadata["comments"] = core_props.comments
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        if core_props.language:
            metadata["language"] = core_props.language
        if hasattr(core_props, 'content_status') and core_props.content_status:
            metadata["content_status"] = core_props.content_status
            
        # Document statistics
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["section_count"] = len(doc.sections)
        metadata["page_count"] = len(doc.sections)  # Approximate
        
        # Table information
        tables = doc.tables
        metadata["table_count"] = len(tables)
        
        # Style information
        styles = doc.styles
        metadata["style_count"] = len(styles)
        
        # Header/footer information
        has_header = False
        has_footer = False
        for section in doc.sections:
            if section.header.is_linked_to_previous == False:
                has_header = True
            if section.footer.is_linked_to_previous == False:
                has_footer = True
        
        metadata["has_headers"] = has_header
        metadata["has_footers"] = has_footer
        
        # Check for images
        try:
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            metadata["image_count"] = image_count
        except Exception as e:
            print(f"Error counting images: {e}")
            
        return convert_metadata(metadata)
    
    except Exception as e:
        print(f"Failed to extract DOCX metadata: {e}")
        return {"Error": "Failed to extract DOCX metadata"}

def extract_video_metadata(file_path):
    """
    Extract detailed metadata from video files using pymediainfo.
    Includes video, audio, and general properties.
    """
    metadata = {}
    try:
        metadata["file_name"] = os.path.basename(file_path)
        metadata["file_size"] = os.path.getsize(file_path)
        
        media_info = MediaInfo.parse(file_path)
        
        # General information
        for track in media_info.tracks:
            if track.track_type == "General":
                general_info = {
                    "Format": track.format,
                    "Format Profile": track.format_profile if hasattr(track, 'format_profile') else None,
                    "Codec": track.codec_id if hasattr(track, 'codec_id') else None,
                    "Duration": f"{track.duration} ms" if hasattr(track, 'duration') else None,
                    "Duration_Formatted": track.other_duration[0] if hasattr(track, 'other_duration') else None,
                    "File Size": f"{track.file_size} bytes" if hasattr(track, 'file_size') else None,
                    "Overall Bitrate": f"{track.overall_bit_rate} bps" if hasattr(track, 'overall_bit_rate') else None,
                    "Created Date": track.file_created_date if hasattr(track, 'file_created_date') else None,
                    "Modified Date": track.file_modified_date if hasattr(track, 'file_modified_date') else None,
                    "Encoded by": track.encoded_by if hasattr(track, 'encoded_by') else None,
                    "Writing Application": track.writing_application if hasattr(track, 'writing_application') else None,
                }
                # Filter out None values
                metadata["General"] = {k: v for k, v in general_info.items() if v is not None}
                
            # Video information
            elif track.track_type == "Video":
                video_info = {
                    "Format": track.format,
                    "Format Profile": track.format_profile if hasattr(track, 'format_profile') else None,
                    "Codec": track.codec_id if hasattr(track, 'codec_id') else None,
                    "Width": f"{track.width} pixels" if hasattr(track, 'width') else None,
                    "Height": f"{track.height} pixels" if hasattr(track, 'height') else None,
                    "Display Aspect Ratio": track.display_aspect_ratio if hasattr(track, 'display_aspect_ratio') else None,
                    "Frame Rate": f"{track.frame_rate} fps" if hasattr(track, 'frame_rate') else None,
                    "Bit Depth": f"{track.bit_depth} bits" if hasattr(track, 'bit_depth') else None,
                    "Bit Rate": f"{track.bit_rate} bps" if hasattr(track, 'bit_rate') else None,
                    "Scan Type": track.scan_type if hasattr(track, 'scan_type') else None,
                    "Color Space": track.color_space if hasattr(track, 'color_space') else None,
                    "Chroma Subsampling": track.chroma_subsampling if hasattr(track, 'chroma_subsampling') else None,
                }
                # Filter out None values
                metadata["Video"] = {k: v for k, v in video_info.items() if v is not None}
                
            # Audio information
            elif track.track_type == "Audio":
                audio_info = {
                    "Format": track.format,
                    "Format Profile": track.format_profile if hasattr(track, 'format_profile') else None,
                    "Codec": track.codec_id if hasattr(track, 'codec_id') else None,
                    "Channel(s)": track.channel_s if hasattr(track, 'channel_s') else None,
                    "Sampling Rate": f"{track.sampling_rate} Hz" if hasattr(track, 'sampling_rate') else None,
                    "Bit Rate": f"{track.bit_rate} bps" if hasattr(track, 'bit_rate') else None,
                    "Compression Mode": track.compression_mode if hasattr(track, 'compression_mode') else None,
                    "Language": track.language if hasattr(track, 'language') else None,
                    "Default": track.default if hasattr(track, 'default') else None,
                }
                # Filter out None values
                if "Audio" not in metadata:
                    metadata["Audio"] = []
                metadata["Audio"].append({k: v for k, v in audio_info.items() if v is not None})
                
            # Subtitle information
            elif track.track_type == "Text":
                text_info = {
                    "Format": track.format,
                    "Codec ID": track.codec_id if hasattr(track, 'codec_id') else None,
                    "Language": track.language if hasattr(track, 'language') else None,
                    "Default": track.default if hasattr(track, 'default') else None,
                    "Forced": track.forced if hasattr(track, 'forced') else None,
                }
                # Filter out None values
                if "Subtitles" not in metadata:
                    metadata["Subtitles"] = []
                metadata["Subtitles"].append({k: v for k, v in text_info.items() if v is not None})
        
        return convert_metadata(metadata)
    
    except Exception as e:
        print(f"Failed to extract video metadata: {e}")
        return {"Error": "Failed to extract video metadata"}

def extract_metadata(file_path):
    """
    Determine file type and extract appropriate metadata.
    Returns a dictionary with detailed metadata based on file type.
    """
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return {"Error": "File not found"}
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # Process based on file extension
    if ext in [".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"]:
        return extract_image_metadata(file_path)
    elif ext == ".pdf":
        return extract_pdf_metadata(file_path)
    elif ext == ".docx":
        return extract_docx_metadata(file_path)
    elif ext in [".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm"]:
        return extract_video_metadata(file_path)
    else:
        print(f"Unsupported file type: {ext}")
        return {"Error": "Unsupported file type"}

def main():
    """
    Main function: demonstrate usage with a sample file.
    """
    if len(sys.argv) != 2:
        print("Usage: python metadata_extractor.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    metadata = extract_metadata(file_path)
    
    # Check if there was an error
    if "Error" in metadata and len(metadata) == 1:
        print(f"Error: {metadata['Error']}")
        sys.exit(1)
    
    # Print formatted JSON output
    print(json.dumps(metadata, indent=2))

if __name__ == "__main__":
    main()